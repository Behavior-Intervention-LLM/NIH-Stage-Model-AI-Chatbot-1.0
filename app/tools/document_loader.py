"""
： PDF、DOC、DOCX 
"""
import os
from typing import List, Dict, Optional
from pathlib import Path
import hashlib


class DocumentChunk:
    """"""
    def __init__(self, content: str, source: str, chunk_index: int, metadata: Optional[Dict] = None):
        self.content = content
        self.source = source
        self.chunk_index = chunk_index
        self.metadata = metadata or {}
        self.id = self._generate_id()
    
    def _generate_id(self) -> str:
        """ID"""
        content_hash = hashlib.md5(
            f"{self.source}_{self.chunk_index}_{self.content[:100]}".encode()
        ).hexdigest()
        return f"{self.source}_{self.chunk_index}_{content_hash[:8]}"
    
    def to_dict(self) -> Dict:
        """"""
        return {
            "id": self.id,
            "content": self.content,
            "source": self.source,
            "chunk_index": self.chunk_index,
            "metadata": self.metadata
        }


class DocumentLoader:
    """"""
    
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        """
        Args:
            chunk_size: （）
            chunk_overlap: 
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def load_pdf(self, file_path: str) -> List[DocumentChunk]:
        """ PDF """
        try:
            import PyPDF2
        except ImportError:
            try:
                import pdfplumber
            except ImportError:
                raise ImportError(" PyPDF2  pdfplumber: pip install PyPDF2 pdfplumber")
        
        chunks = []
        source_name = os.path.basename(file_path)
        
        try:
            #  pdfplumber（）
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                full_text = ""
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        full_text += f"\n\n---  {page_num}  ---\n\n{text}"
                
                if full_text:
                    chunks = self._split_text(full_text, source_name)
        except ImportError:
            #  PyPDF2
            import PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                full_text = ""
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text = page.extract_text()
                    if text:
                        full_text += f"\n\n---  {page_num}  ---\n\n{text}"
                
                if full_text:
                    chunks = self._split_text(full_text, source_name)
        
        return chunks
    
    def load_docx(self, file_path: str) -> List[DocumentChunk]:
        """ DOCX """
        try:
            from docx import Document
        except ImportError:
            raise ImportError(" python-docx: pip install python-docx")
        
        doc = Document(file_path)
        source_name = os.path.basename(file_path)
        
        full_text = ""
        for para in doc.paragraphs:
            if para.text.strip():
                full_text += para.text + "\n\n"
        
        # 
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    full_text += row_text + "\n"
            full_text += "\n"
        
        chunks = self._split_text(full_text, source_name) if full_text else []
        return chunks
    
    def load_doc(self, file_path: str) -> List[DocumentChunk]:
        """ DOC """
        #  python-docx（）
        try:
            return self.load_docx(file_path)
        except:
            #  docx2txt
            try:
                import docx2txt
                source_name = os.path.basename(file_path)
                text = docx2txt.process(file_path)
                chunks = self._split_text(text, source_name) if text else []
                return chunks
            except ImportError:
                raise ImportError(" docx2txt: pip install docx2txt")
    
    def load_file(self, file_path: str) -> List[DocumentChunk]:
        """"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f": {file_path}")
        
        ext = file_path.suffix.lower()
        
        if ext == '.pdf':
            return self.load_pdf(str(file_path))
        elif ext == '.docx':
            return self.load_docx(str(file_path))
        elif ext == '.doc':
            return self.load_doc(str(file_path))
        else:
            raise ValueError(f": {ext}")
    
    def _split_text(self, text: str, source: str) -> List[DocumentChunk]:
        """"""
        chunks = []
        start = 0
        chunk_index = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # 
            if end < len(text):
                # 、、
                for i in range(end, max(start + self.chunk_size - 100, start), -1):
                    if text[i] in '。！？\n':
                        end = i + 1
                        break
            
            chunk_text = text[start:end].strip()
            if chunk_text:
                chunks.append(DocumentChunk(
                    content=chunk_text,
                    source=source,
                    chunk_index=chunk_index,
                    metadata={"start": start, "end": end}
                ))
                chunk_index += 1
            
            # （）
            start = end - self.chunk_overlap
        
        return chunks
    
    def load_directory(self, directory: str) -> List[DocumentChunk]:
        """"""
        directory = Path(directory)
        all_chunks = []
        
        supported_extensions = ['.pdf', '.doc', '.docx']
        
        for ext in supported_extensions:
            for file_path in directory.glob(f'*{ext}'):
                try:
                    chunks = self.load_file(str(file_path))
                    all_chunks.extend(chunks)
                    print(f"✓ : {file_path.name} ({len(chunks)} )")
                except Exception as e:
                    print(f"✗  {file_path.name}: {e}")
        
        return all_chunks
